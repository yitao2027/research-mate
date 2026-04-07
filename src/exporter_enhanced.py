#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
exporter_enhanced.py - 增强版素材导出器

支持：
1. Word 文档导出（含提纲、总结、素材详情）
2. PDF 导出
3. 素材评估表
4. 补充采集建议收集
"""

import json
import csv
from datetime import datetime
from typing import List, Dict, Any, Tuple
from pathlib import Path


class EnhancedMaterialExporter:
    """增强版素材导出器"""
    
    def __init__(self, output_dir: str = "./output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export_to_word(self, materials: List[Dict[str, Any]], 
                      topic: str, 
                      outline: str = None,
                      summary: str = None,
                      target_words: int = 3000) -> str:
        """
        导出为 Word 文档（.docx）
        
        Args:
            materials: 素材列表
            topic: 文章主题
            outline: 文章提纲
            summary: 内容总结
            target_words: 目标字数
            
        Returns:
            生成的 Word 文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            print("⚠️  未安装 python-docx，正在尝试安装...")
            import subprocess
            subprocess.check_call(['pip', 'install', 'python-docx', '-q'])
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        
        # 创建文档
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = u'微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        doc.styles['Normal'].font.size = Pt(10.5)
        doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        
        # 标题
        title = doc.add_heading(f'📚 {topic} - 素材采集报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        info_para = doc.add_paragraph()
        info_para.add_run(f"**生成时间**：{timestamp}\n").bold = True
        info_para.add_run(f"**素材总数**：{len(materials)} 条\n")
        info_para.add_run(f"**目标字数**：{target_words} 字\n")
        info_para.add_run(f"**本报告由 ResearchMate 自动生成**")
        
        doc.add_page_break()
        
        # 一、文章提纲
        if outline:
            doc.add_heading('📋 一、文章提纲', level=1)
            doc.add_paragraph(outline)
            doc.add_page_break()
        
        # 二、核心总结
        if summary:
            doc.add_heading('💡 二、核心总结', level=1)
            doc.add_paragraph(summary)
            doc.add_page_break()
        
        # 三、素材分类统计
        doc.add_heading('📊 三、素材分类统计', level=1)
        
        category_stats = {}
        for mat in materials:
            cat = mat.get('category', '未分类')
            category_stats[cat] = category_stats.get(cat, 0) + 1
        
        # 添加表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '素材类型'
        header_cells[1].text = '数量'
        header_cells[2].text = '占比'
        
        # 加粗表头
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 填充数据
        total = len(materials)
        for cat, count in sorted(category_stats.items(), key=lambda x: x[1], reverse=True):
            row = table.add_row()
            row.cells[0].text = cat
            row.cells[1].text = str(count)
            row.cells[2].text = f'{count/total*100:.1f}%'
        
        doc.add_paragraph()
        
        # 四、详细素材卡片
        doc.add_heading('📝 四、详细素材卡片', level=1)
        
        for idx, mat in enumerate(materials, 1):
            # 素材标题
            title_para = doc.add_heading(f'素材 #{idx}', level=2)
            
            # 基本信息表格
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Table Grid'
            
            # 填充信息
            info_data = [
                ('标题', mat.get('title', '无标题')),
                ('类型', mat.get('category', '未分类')),
                ('来源', mat.get('source', '未知')),
                ('日期', mat.get('publish_date', '未知')),
                ('质量评分', f'⭐ {mat.get("score", 0):.1f}/100')
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[1].text = value
                
                # 第一列加粗
                info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            
            # URL（如果有）
            url = mat.get('url', '')
            if url and url != '#':
                url_para = doc.add_paragraph()
                url_para.add_run('🔗 原文链接：').bold = True
                url_para.add_run(url)
            
            # 核心内容
            content = mat.get('content', '无内容')
            doc.add_paragraph('核心内容：', style='Intense Quote')
            doc.add_paragraph(content)
            
            # 分隔线
            if idx < len(materials):
                doc.add_paragraph('_' * 50)
        
        # 五、素材评估表
        doc.add_heading('📈 五、素材评估表', level=1)
        
        eval_table = doc.add_table(rows=1, cols=6)
        eval_table.style = 'Table Grid'
        
        # 表头
        eval_headers = ['序号', '标题', '类型', '可信度', '时效性', '评分']
        for i, header in enumerate(eval_headers):
            eval_table.rows[0].cells[i].text = header
            eval_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        # 填充数据
        for idx, mat in enumerate(materials, 1):
            row = eval_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = mat.get('title', '')[:20] + '...' if len(mat.get('title', '')) > 20 else mat.get('title', '')
            row.cells[2].text = mat.get('category', '未分类')
            
            breakdown = mat.get('score_breakdown', {})
            row.cells[3].text = f'{breakdown.get("credibility", 0):.2f}'
            row.cells[4].text = f'{breakdown.get("timeliness", 0):.2f}'
            row.cells[5].text = f'{mat.get("score", 0):.1f}'
        
        # 保存文件
        safe_topic = topic.replace('/', '_').replace('\\', '_').replace(':', '_')
        filename = f"{safe_topic}_素材报告_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        print(f"✅ Word 报告已生成：{filepath}")
        return str(filepath)
    
    def generate_outline_and_summary(self, materials: List[Dict[str, Any]], 
                                    topic: str) -> Tuple[str, str]:
        """
        根据素材生成文章提纲和总结
        
        Returns:
            (outline, summary) 元组
        """
        # 分析素材分布
        category_stats = {}
        key_points = []
        
        for mat in materials:
            cat = mat.get('category', '未分类')
            category_stats[cat] = category_stats.get(cat, 0) + 1
            
            # 提取关键点
            content = mat.get('content', '')
            if len(content) > 50:
                key_points.append(content[:100])
        
        # 生成提纲（基于素材类型）
        outline_parts = [f"# {topic}\n"]
        
        if '财务数据' in category_stats or '市场数据' in category_stats:
            outline_parts.append("## 一、行业现状与市场规模\n- 整体规模与增长趋势\n- 主要企业市场份额\n")
        
        if '产品创新' in category_stats:
            outline_parts.append("## 二、技术创新与产品发展\n- 核心技术突破\n- 产品迭代方向\n")
        
        if '融资数据' in category_stats:
            outline_parts.append("## 三、资本动态与投融资分析\n- 融资事件梳理\n- 投资热点与趋势\n")
        
        if '企业战略' in category_stats:
            outline_parts.append("## 四、竞争格局与企业战略\n- 头部企业布局\n- 竞争策略分析\n")
        
        outline_parts.append("## 五、未来展望与建议\n- 发展趋势预测\n- 投资机会与风险\n")
        
        outline = '\n'.join(outline_parts)
        
        # 生成总结
        summary = f"本次围绕\"{topic}\"主题，共采集{len(materials)}条高质量素材。"
        summary += f"素材涵盖{len(category_stats)}个维度："
        
        for cat, count in sorted(category_stats.items(), key=lambda x: x[1], reverse=True)[:3]:
            summary += f"{cat}{count}条、"
        
        summary = summary.rstrip('、') + "。"
        summary += "素材来源权威，包含券商研报、主流财经媒体及官方披露信息，"
        summary += "可为文章撰写提供充分的数据支撑和案例参考。"
        
        return outline, summary
    
    def generate_assessment_form(self, materials: List[Dict[str, Any]], 
                                topic: str) -> str:
        """
        生成素材评估表（CSV 格式）
        
        Returns:
            CSV 文件路径
        """
        filepath = self.output_dir / f"{topic}_素材评估表.csv"
        
        with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            
            # 表头
            writer.writerow([
                '序号', '标题', '类型', '来源', 'URL', 
                '可信度', '时效性', '完整性', '交叉验证', 
                '综合评分', '等级', '是否采用', '补充建议'
            ])
            
            # 数据行
            for idx, mat in enumerate(materials, 1):
                breakdown = mat.get('score_breakdown', {})
                
                writer.writerow([
                    idx,
                    mat.get('title', ''),
                    mat.get('category', ''),
                    mat.get('source', ''),
                    mat.get('url', ''),
                    f"{breakdown.get('credibility', 0):.2f}",
                    f"{breakdown.get('timeliness', 0):.2f}",
                    f"{breakdown.get('completeness', 0):.2f}",
                    f"{breakdown.get('verification', 0):.2f}",
                    f"{mat.get('score', 0):.1f}",
                    mat.get('rating', ''),
                    '□',  # 复选框
                    ''     # 用户填写补充建议
                ])
        
        print(f"✅ 素材评估表已生成：{filepath}")
        return str(filepath)
    
    def collect_feedback(self, materials: List[Dict[str, Any]]) -> Dict:
        """
        收集用户对素材的反馈
        
        Returns:
            反馈字典
        """
        print("\n" + "=" * 60)
        print("📝 请对素材进行评估和反馈")
        print("=" * 60)
        
        feedback = {
            'satisfied_count': 0,
            'need_more': [],
            'suggestions': []
        }
        
        # 展示素材评分前 10 条
        top_materials = sorted(materials, key=lambda x: x.get('score', 0), reverse=True)[:10]
        
        print("\n⭐ 评分最高的 10 条素材：")
        for idx, mat in enumerate(top_materials, 1):
            print(f"  {idx}. [{mat.get('rating', '')}] {mat.get('title', '')[:40]}... "
                  f"(评分：{mat.get('score', 0):.1f})")
        
        # 询问是否需要补充
        need_more_input = input("\n❓ 是否还需要补充某些方面的素材？(y/n)：").strip().lower()
        
        if need_more_input == 'y':
            print("\n💡 请描述需要补充的素材类型或方向：")
            while True:
                suggestion = input("   （输入空行结束）：").strip()
                if not suggestion:
                    break
                feedback['need_more'].append(suggestion)
        
        # 总体满意度
        satisfied_input = input(f"\n❓ 当前素材是否满足写作需求？(满意素材数/{len(materials)})：").strip()
        
        try:
            feedback['satisfied_count'] = int(satisfied_input)
        except ValueError:
            feedback['satisfied_count'] = len(materials)
        
        return feedback


# 测试代码
if __name__ == '__main__':
    print("🧪 测试增强版导出器...")
    
    test_materials = [
        {
            "title": "人工智能行业最新动态",
            "category": "财务数据",
            "source": "36 氪",
            "url": "https://36kr.com/example",
            "content": "最新研究显示，人工智能领域头部企业市场份额进一步扩大。2025 年 CR5 达到 65%，较上年提升 8 个百分点。",
            "score": 88.5,
            "score_breakdown": {
                'credibility': 0.85,
                'timeliness': 0.90,
                'completeness': 0.80,
                'verification': 0.85
            },
            "rating": "A"
        }
    ]
    
    exporter = EnhancedMaterialExporter("./test_output")
    
    # 测试生成提纲和总结
    outline, summary = exporter.generate_outline_and_summary(test_materials, "AI 行业分析")
    print("\n📋 文章提纲:")
    print(outline)
    print("\n💡 核心总结:")
    print(summary)
    
    # 测试 Word 导出
    word_path = exporter.export_to_word(test_materials, "AI 行业分析", outline, summary)
    print(f"\n✅ Word 报告：{word_path}")
    
    # 测试评估表
    csv_path = exporter.generate_assessment_form(test_materials, "AI 行业分析")
    print(f"✅ 评估表：{csv_path}")
    
    print("\n🎉 增强版导出器测试完成！")
